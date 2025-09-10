import argparse
import json
import logging
import os
import re
import sys
from datetime import datetime
from html import escape
from pathlib import Path
from typing import Dict, Any, List
import pandas as pd
from dotenv import load_dotenv
try:
    from docx import Document  # provided by python-docx
except (ImportError, AttributeError) as exc:
    raise ImportError(
        "Failed to import Document from 'docx'. Install the correct package with 'pip install --upgrade python-docx' "
        "and uninstall any conflicting 'docx' package via 'pip uninstall docx'."
    ) from exc
from unidecode import unidecode

#!/usr/bin/env python3
"""
Generate offer sheets for leads without a proper website (empty website field or containing a Facebook link).

Inputs (default filenames expected in current working directory):
    - Lead-Auto-Ergebnis.xlsx              (Excel sheet with lead data)
    - Angebot-Webseitenservice.docx        (DOCX template containing placeholders like {{COMPANY}}, {{WEBSITE}}, etc.)

Output:
    - One folder per company at: offer-sheets/{Company_name_sanitized}/
    - Inside each folder: A customized DOCX offer sheet (Angebot-Webseitenservice-{Company}.docx)
    - (Optional) a JSON metadata dump for traceability

Filtering rule:
    Select rows where the (case-insensitive) column named one of ["Website", "website", "Webseite", "webseite", "URL", "Url"]
    is either empty / NaN OR contains the substring "facebook" (e.g. facebook.com pages instead of a real site).

Placeholder logic:
    For every column in the Excel row, a placeholder {{COLUMN_NAME_NORMALIZED}} will be replaced.
    Normalization: uppercase, spaces -> underscores, umlauts handled, non-alnum removed.
    Example: Column "Company Name" -> placeholder {{COMPANY_NAME}}
    If the template contains placeholders not present in data, they are left untouched.

Install dependencies (if not already):
    pip install pandas python-docx openpyxl unidecode

Run:
    python lead_filter_pipeline.py
    (or adjust CLI args: -e path/to/excel.xlsx -t path/to/template.docx -o offer-sheets)

"""



# ---------------- Configuration ---------------- #

# New: resolve paths relative to monorepo root (this file lives in Backend/)
ROOT_DIR = Path(__file__).resolve().parents[1]

# Defaults now point to new locations
DEFAULT_EXCEL = str(ROOT_DIR / "Backend" / "data" / "Leads-Auto-Ergebnis.xlsx")
DEFAULT_TEMPLATE = str(ROOT_DIR / "templates" / "docx" / "Angebot-Webseitenservice.docx")
DEFAULT_OUTPUT_DIR = str(ROOT_DIR / "Backend" / "offer-sheets")

WEBSITE_COLUMN_CANDIDATES = ["website", "webseite", "url"]

LOG_FORMAT = "%(levelname)s %(message)s"


# ---------------- Utility Functions ---------------- #

def setup_logging(verbose: bool):
        logging.basicConfig(
                level=logging.DEBUG if verbose else logging.INFO,
                format=LOG_FORMAT
        )


def find_website_column(df: pd.DataFrame) -> str:
        cols_lower = {c.lower(): c for c in df.columns}
        for candidate in WEBSITE_COLUMN_CANDIDATES:
                if candidate in cols_lower:
                        return cols_lower[candidate]
        # fallback: heuristic search
        for c in df.columns:
                if "web" in c.lower() or "site" in c.lower():
                        return c
        raise ValueError("Could not determine website column. Please rename one column to 'Website'.")


def slugify(text: str) -> str:
        if not isinstance(text, str):
                text = str(text) if text is not None else "unknown"
        text = unidecode(text)
        text = text.strip()
        text = re.sub(r"[^\w\s-]", "", text)
        text = re.sub(r"[\s_-]+", "-", text)
        text = re.sub(r"^-+|-+$", "", text)
        return text or "company"


def normalize_column_name(col: str) -> str:
        col = unidecode(str(col))
        col = col.upper()
        col = re.sub(r"[^A-Z0-9]+", "_", col)
        col = re.sub(r"_+", "_", col).strip("_")
        return col


def build_placeholder_map(row: pd.Series) -> Dict[str, str]:
        placeholders = {}
        for col, val in row.items():
                key = normalize_column_name(col)  # NORMALIZED_KEY (e.g. IHR_NAME)
                value = "" if pd.isna(val) else str(val)

                # Prepare a set of candidate placeholder strings to match common template styles
                variants = set()

                # Normalized forms (underscore-style)
                variants.update({f"{{{{{key}}}}}", f"{{{key}}}", f"[{key}]"})

                # Human-friendly forms based on original column header
                col_str = str(col).strip()
                col_nod = unidecode(col_str)

                # Add several casing variants for the original header and the diacritics-free version
                for base in {col_str, col_nod}:
                        variants.update({
                                f"[{base}]",
                                f"[{base.lower()}]",
                                f"[{base.upper()}]",
                                f"[{base.title()}]",
                                f"{{{base}}}",
                                f"{{{{{base}}}}}",
                        })

                # Also allow space vs underscore variants (e.g. IHR_NAME <-> 'Ihr Name')
                space_key = key.replace("_", " ")
                for base in {key, space_key}:
                        variants.update({
                                f"[{base}]",
                                f"{{{base}}}",
                                f"{{{{{base}}}}}",
                                f"[{base.lower()}]",
                                f"[{base.upper()}]",
                                f"[{base.title()}]",
                        })

                # Write all variants into the map
                for v in variants:
                        placeholders[v] = value

        # Add some convenience/computed placeholders (date) in common styles
        date_val = datetime.now().strftime("%d.%m.%Y")
        for dkey in ["{{DATE}}", "{DATE}", "[DATE]", "{{Date}}", "{Date}", "[Date]"]:
                if dkey not in placeholders:
                        placeholders[dkey] = date_val
        return placeholders


def replace_placeholders_in_doc(doc: Document, mapping: Dict[str, str]):
        # Replace in main document paragraphs
        for p in doc.paragraphs:
                replace_in_run_collection(p.runs, mapping)

        # Replace in main document tables
        for table in doc.tables:
                for row in table.rows:
                        for cell in row.cells:
                                for p in cell.paragraphs:
                                        replace_in_run_collection(p.runs, mapping)

        # Replace in headers and footers of each section
        for section in doc.sections:
                header = section.header
                footer = section.footer
                for p in header.paragraphs:
                        replace_in_run_collection(p.runs, mapping)
                for table in header.tables:
                        for row in table.rows:
                                for cell in row.cells:
                                        for p in cell.paragraphs:
                                                replace_in_run_collection(p.runs, mapping)

                for p in footer.paragraphs:
                        replace_in_run_collection(p.runs, mapping)
                for table in footer.tables:
                        for row in table.rows:
                                for cell in row.cells:
                                        for p in cell.paragraphs:
                                                replace_in_run_collection(p.runs, mapping)


def replace_in_run_collection(runs, mapping: Dict[str, str]):
        # Approach: join full text, replace, then rebuild runs minimally
        full_text = "".join(run.text for run in runs)
        original = full_text
        for placeholder, value in mapping.items():
                if placeholder in full_text:
                        full_text = full_text.replace(placeholder, value)
        # Remove any leftover placeholder patterns like {{...}}, {...}, or [...] to avoid leaking tokens
        full_text = re.sub(r"\{\{[^}]+\}\}|\{[^}]+\}|\[[^\]]+\]", "", full_text)

        if full_text != original:
                # Clear runs and set first run to new text
                if runs:
                        runs[0].text = full_text
                        for r in runs[1:]:
                                r.text = ""


# --- New helpers for text/markdown templates ---
def load_text_template(template_path: Path) -> str:
        try:
                return template_path.read_text(encoding="utf-8")
        except Exception as e:
                logging.debug(f"Failed to read text template {template_path}: {e}")
                return ""


def replace_placeholders_in_text(text: str, mapping: Dict[str, str]) -> str:
        # Simple replace loop - mapping contains many common variants
        for placeholder, value in mapping.items():
                if placeholder in text:
                        text = text.replace(placeholder, value)
        # Remove any leftover placeholder patterns
        text = re.sub(r"\{\{[^}]+\}\}|\{[^}]+\}|\[[^\]]+\]", "", text)
        return text


def generate_text_templates_for_offer(row: pd.Series, mapping: Dict[str, str], target_dir: Path,
                                      email_template: Path, phone_template: Path, overwrite: bool = False):
        # Use monorepo root for resolving templates
        templates_root = ROOT_DIR

        # Resolve templates relative to repo root if not absolute
        email_tpl = email_template if email_template.is_absolute() else (templates_root / email_template)
        phone_tpl = phone_template if phone_template.is_absolute() else (templates_root / phone_template)

        for tpl_path, out_name in ((email_tpl, "cold_email.md"), (phone_tpl, "cold_phone_call.md")):
                if not tpl_path.exists():
                        logging.debug(f"Text template not found: {tpl_path} - skipping {out_name}")
                        continue
                out_file = target_dir / out_name
                if out_file.exists() and not overwrite:
                        logging.info(f"Skipping existing {out_name} for {target_dir.name}")
                        continue
                txt = load_text_template(tpl_path)
                if not txt:
                        continue
                filled = replace_placeholders_in_text(txt, mapping)
                try:
                        out_file.write_text(filled, encoding="utf-8")
                        logging.debug(f"Wrote text template: {out_file}")
                except Exception as e:
                        logging.error(f"Failed writing {out_file}: {e}")


def load_excel(excel_path: Path) -> pd.DataFrame:
        df = pd.read_excel(excel_path)
        if df.empty:
                logging.warning("Excel file is empty.")
        return df


def filter_rows(df: pd.DataFrame, website_col: str) -> pd.DataFrame:
        col = df[website_col]
        # Standardize to string where possible
        col_str = col.astype(str).str.strip().replace({"nan": ""})
        mask_empty = col_str.eq("") | col.isna()
        mask_fb = col_str.str.contains("facebook", case=False, na=False)
        filtered = df[mask_empty | mask_fb].copy()
        logging.info(f"Filtered {len(filtered)}/{len(df)} rows needing website offer.")
        return filtered


def ensure_dir(path: Path):
        path.mkdir(parents=True, exist_ok=True)


def _row_value_by_keys(row: pd.Series, key_variants: List[str]) -> str:
        col_map = {normalize_column_name(c): c for c in row.index}
        for k in key_variants:
                if k in col_map:
                        v = row[col_map[k]]
                        if pd.isna(v):
                                return ""
                        return str(v)
        return ""


def _sanitize_phone_for_tel(phone: str) -> str:
        if not phone:
            return ""
        return re.sub(r"[^0-9+]+", "", phone)


def _ensure_http(url: str) -> str:
        if not url:
                return ""
        u = url.strip()
        if not re.match(r"^https?://", u, flags=re.IGNORECASE):
                return "http://" + u
        return u


def generate_lead_html(row: pd.Series, company_name: str, target_dir: Path, mapping: Dict[str, str], overwrite: bool = False) -> Path:
        """Render lead HTML summary from an external template (templates/html/lead_summary_template.html)."""
        phone = _row_value_by_keys(row, [
                "PHONE", "TELEFON", "TEL", "MOBILE", "HANDY", "PHONE_NUMBER"
        ])
        email = _row_value_by_keys(row, [
                "EMAIL", "E_MAIL", "MAIL", "EMAIL_ADDRESS"
        ])
        website = _row_value_by_keys(row, [
                "WEBSITE", "WEBSEITE", "URL"
        ])
        city = _row_value_by_keys(row, [
                "CITY", "STADT", "ORT"
        ])
        industry = _row_value_by_keys(row, [
                "INDUSTRY", "BRANCHE", "BUSINESS_TYPE"
        ])
        contact = _row_value_by_keys(row, [
                "ANSPRECHPARTNER", "CONTACT", "CONTACT_NAME", "IHR_NAME", "OWNER", "MANAGER", "NAME"
        ])

        phone_href = _sanitize_phone_for_tel(phone)
        email_href = email.strip() if email else ""
        website_href = _ensure_http(website)

        # Escape for HTML display
        comp_e = escape(company_name)
        phone_e = escape(phone or "")
        email_e = escape(email or "")
        website_e = escape(website or "")
        city_e = escape(city or "")
        industry_e = escape(industry or "")
        contact_e = escape(contact or "")

        # Attempt to include generated markdown templates (email / phone) as drawers
        email_md = ""
        phone_md = ""
        try:
                email_path = target_dir / "cold_email.md"
                phone_path = target_dir / "cold_phone_call.md"
                if email_path.exists():
                        email_md = escape(email_path.read_text(encoding="utf-8"))
                if phone_path.exists():
                        phone_md = escape(phone_path.read_text(encoding="utf-8"))
        except Exception:
                # non-fatal: leave strings empty
                email_md = email_md or ""
                phone_md = phone_md or ""

        details_items = []
        if contact_e:
                details_items.append(f"<li><span class=\"font-medium\">Contact:</span> {contact_e}</li>")
        if city_e:
                details_items.append(f"<li><span class=\"font-medium\">City:</span> {city_e}</li>")
        if industry_e:
                details_items.append(f"<li><span class=\"font-medium\">Industry:</span> {industry_e}</li>")

        generated_at = datetime.now().strftime("%d.%m.%Y %H:%M")

        # Build dynamic HTML blocks for the template
        phone_block = (
            f"<div><span class=\"text-sm text-gray-500 dark:text-gray-400\">Phone:</span> "
            + (f"<a class=\"text-blue-600 dark:text-blue-400 hover:underline\" href=\"tel:{phone_href}\">{phone_e}</a>" if phone_href else phone_e)
            + "</div>"
        ) if (phone_e or phone_href) else ""

        email_block = (
            f"<div><span class=\"text-sm text-gray-500 dark:text-gray-400\">Email:</span> "
            + (f"<a class=\"text-blue-600 dark:text-blue-400 hover:underline\" href=\"mailto:{email_href}\">{email_e}</a>" if email_href else email_e)
            + "</div>"
        ) if (email_e or email_href) else ""

        website_block = (
            f"<div><span class=\"text-sm text-gray-500 dark:text-gray-400\">Website:</span> "
            + (f"<a class=\"text-blue-600 dark:text-blue-400 hover:underline\" target=\"_blank\" rel=\"noopener noreferrer\" href=\"{website_href}\">{website_e}</a>" if website_href else website_e)
            + "</div>"
        ) if (website_e or website_href) else ""

        details_list = ("<ul class=\"mt-6 space-y-1 text-sm text-gray-700 dark:text-gray-200\">" + "".join(details_items) + "</ul>") if details_items else ""

        email_md_block = (
            "<details class=\"bg-gray-50 dark:bg-gray-900 border border-gray-100 dark:border-gray-800 rounded-lg p-4\">\n"
            "  <summary class=\"cursor-pointer font-medium text-gray-900 dark:text-gray-100\">Cold email (click to expand)</summary>\n"
            f"  <pre class=\"mt-3 whitespace-pre-wrap bg-white dark:bg-gray-800 p-3 rounded text-sm text-gray-800 dark:text-gray-100 font-mono overflow-x-auto\">{email_md}</pre>\n"
            "</details>"
        ) if email_md else ""

        phone_md_block = (
            "<details class=\"bg-gray-50 dark:bg-gray-900 border border-gray-100 dark:border-gray-800 rounded-lg p-4\">\n"
            "  <summary class=\"cursor-pointer font-medium text-gray-900 dark:text-gray-100\">Cold phone script (click to expand)</summary>\n"
            f"  <pre class=\"mt-3 whitespace-pre-wrap bg-white dark:bg-gray-800 p-3 rounded text-sm text-gray-800 dark:text-gray-100 font-mono overflow-x-auto\">{phone_md}</pre>\n"
            "</details>"
        ) if phone_md else ""

        # Load the external HTML template
        tpl_path = ROOT_DIR / "templates" / "html" / "lead_summary_template.html"
        html_tpl = load_text_template(tpl_path)
        if not html_tpl:
                logging.error(f"HTML template not found: {tpl_path}")
                return target_dir / "lead_summary.html"

        # Extend mapping with blocks and common keys
        mapping_local = dict(mapping)
        mapping_local.update({
                "{{BusinessName}}": comp_e,
                "{BusinessName}": comp_e,
                "{{COMPANY_NAME}}": comp_e,
                "{COMPANY_NAME}": comp_e,
                "{{PHONE_BLOCK}}": phone_block,
                "{PHONE_BLOCK}": phone_block,
                "{{EMAIL_BLOCK}}": email_block,
                "{EMAIL_BLOCK}": email_block,
                "{{WEBSITE_BLOCK}}": website_block,
                "{WEBSITE_BLOCK}": website_block,
                "{{DETAILS_LIST}}": details_list,
                "{DETAILS_LIST}": details_list,
                "{{EMAIL_MD_BLOCK}}": email_md_block,
                "{EMAIL_MD_BLOCK}": email_md_block,
                "{{PHONE_MD_BLOCK}}": phone_md_block,
                "{PHONE_MD_BLOCK}": phone_md_block,
                "{{GENERATED_AT}}": escape(generated_at),
                "{GENERATED_AT}": escape(generated_at),
        })

        filled = replace_placeholders_in_text(html_tpl, mapping_local)

        out_path = target_dir / "lead_summary.html"
        if out_path.exists() and not overwrite:
                logging.info(f"Skipping existing HTML summary for {target_dir.name}")
                return out_path
        try:
                out_path.write_text(filled, encoding="utf-8")
                logging.debug(f"Wrote HTML summary: {out_path}")
        except Exception as e:
                logging.error(f"Failed writing HTML summary {out_path}: {e}")
        return out_path


def _register_placeholder_variants(mapping: Dict[str, str], key: str, value: str):
        if value is None:
                value = ""
        # Register the exact key in common placeholder styles
        for k in {key, key.strip()}:
                mapping[f"{{{{{k}}}}}"] = value
                mapping[f"{{{k}}}"] = value
                mapping[f"[{k}]"] = value


def enrich_placeholders_with_env_and_aliases(mapping: Dict[str, str], row: pd.Series, company_name: str):
        # Row-derived fields (prefer data from the sheet)
        city = _row_value_by_keys(row, ["CITY", "STADT", "ORT"]) or os.getenv("CITY", "")
        website = _row_value_by_keys(row, ["WEBSITE", "WEBSEITE", "URL"]) or os.getenv("YOUR_WEBSITE", "")
        phone = _row_value_by_keys(row, ["PHONE", "TELEFON", "TEL", "MOBILE", "HANDY", "PHONE_NUMBER"]) or os.getenv("YOUR_PHONE", "")
        email = _row_value_by_keys(row, ["EMAIL", "E_MAIL", "MAIL", "EMAIL_ADDRESS"]) or os.getenv("YOUR_EMAIL", "")
        industry = _row_value_by_keys(row, ["INDUSTRY", "BRANCHE", "BUSINESS_TYPE"]) or os.getenv("INDUSTRY", "")
        contact = _row_value_by_keys(row, ["ANSPRECHPARTNER", "CONTACT", "CONTACT_NAME", "IHR_NAME", "OWNER", "MANAGER", "NAME"]) or ""
        first_name = contact.split()[0] if contact else ""

        # Static/env-backed fields (sender)
        your_name = os.getenv("YOUR_NAME", "")
        your_title = os.getenv("YOUR_TITLE", "")
        your_company = os.getenv("YOUR_COMPANY", "")
        your_website = os.getenv("YOUR_WEBSITE", website)
        your_phone = os.getenv("YOUR_PHONE", phone)
        your_email = os.getenv("YOUR_EMAIL", email)
        calendar_link = os.getenv("CALENDAR_LINK", "")
        project_link = os.getenv("PROJECT_LINK", "")
        short_outcome = os.getenv("SHORT_OUTCOME", "")
        default_price = os.getenv("DEFAULT_PRICE", "")
        default_pages = os.getenv("DEFAULT_PAGES", "")
        default_timeline = os.getenv("DEFAULT_TIMELINE", "")
        support_period = os.getenv("SUPPORT_PERIOD", "")

        # Lead vs sender aliases for templates (email/phone md)
        alias_values = {
                # Lead/company being contacted
                "BusinessName": company_name,
                "Business Name": company_name,
                "LeadCompany": company_name,
                # Context
                "City": city,
                "FirstName": first_name,
                "Industry": industry,
                # Contact info from sheet (fallback to sender where helpful)
                "Phone": your_phone or phone,
                "Email": email or your_email,
                "Website": your_website or website,
                "URL": website,
                # Offer defaults
                "Price": default_price,
                "Pages": default_pages,
                "Timeline": default_timeline,
                "SupportPeriod": support_period,
                # Social proof / links
                "ProjectLink": project_link,
                "ShortOutcome": short_outcome,
                "CalendarLink": calendar_link,
                "Link": calendar_link,
                "Short URL": calendar_link,
                # Sender signature
                "YourName": your_name,
                "Your Title": your_title,
                "YourTitle": your_title,
                "Your Company": your_company,
                # Note: In email template {{Company}} denotes sender company
                "Company": your_company,
                # Phone script extras
                "Owner/Manager Name": contact,
                "Name": contact,
                "Role": os.getenv("DEFAULT_ROLE", "Owner"),
        }

        for k, v in alias_values.items():
                _register_placeholder_variants(mapping, k, v)


def generate_offer(
        row: pd.Series,
        template_path: Path,
        company_col: str,
        output_root: Path,
        overwrite: bool
) -> Path:
        company_raw = row.get(company_col, "Unknown Company")
        company_name = str(company_raw) if not pd.isna(company_raw) else "Unknown Company"
        company_slug = slugify(company_name)
        target_dir = output_root / company_slug
        ensure_dir(target_dir)

        output_doc = target_dir / f"Angebot-Webseitenservice-{company_slug}.docx"
        meta_json = target_dir / "metadata.json"

        # Build placeholders once from the row
        placeholders = build_placeholder_map(row)
        enrich_placeholders_with_env_and_aliases(placeholders, row, company_name)

        # Generate or skip DOCX based on overwrite flag
        if output_doc.exists() and not overwrite:
                logging.info(f"Skipping existing offer DOCX for {company_name} ({output_doc})")
        else:
                try:
                        doc = Document(str(template_path))
                        replace_placeholders_in_doc(doc, placeholders)
                        doc.save(str(output_doc))
                        logging.debug(f"Saved offer: {output_doc}")
                except Exception as e:
                        logging.error(f"Failed to generate DOCX for {company_name}: {e}")

        # Always (re)write metadata or respect overwrite?
        try:
                if not meta_json.exists() or overwrite:
                        with meta_json.open("w", encoding="utf-8") as f:
                                json.dump(
                                        {
                                                "company": company_name,
                                                "company_slug": company_slug,
                                                "generated_at": datetime.utcnow().isoformat() + "Z",
                                                "placeholders": placeholders
                                        },
                                        f,
                                        ensure_ascii=False,
                                        indent=2
                                )
        except Exception as e:
                logging.error(f"Failed writing metadata for {company_name}: {e}")

        # Generate cold outreach markdown files
        try:
                # Determine language from environment (.env) or fallback to 'en'
                lang = (os.getenv("TEMPLATE_LANG") or os.getenv("LANG") or "en").strip()
                templates_dir = ROOT_DIR / "templates"
                email_tpl = templates_dir / lang / "cold_email_template.md"
                phone_tpl = templates_dir / lang / "cold_phone_call_template.md"

                # Fallback to English templates if localized ones are missing
                en_email = templates_dir / "en" / "cold_email_template.md"
                en_phone = templates_dir / "en" / "cold_phone_call_template.md"
                if not email_tpl.exists():
                        logging.debug(f"Email template for lang '{lang}' not found at {email_tpl}; falling back to {en_email}")
                        email_tpl = en_email
                if not phone_tpl.exists():
                        logging.debug(f"Phone template for lang '{lang}' not found at {phone_tpl}; falling back to {en_phone}")
                        phone_tpl = en_phone

                generate_text_templates_for_offer(
                        row=row,
                        mapping=placeholders,
                        target_dir=target_dir,
                        email_template=email_tpl,
                        phone_template=phone_tpl,
                        overwrite=overwrite
                )
        except Exception as e:
                logging.error(f"Failed generating text templates for {company_name}: {e}")

        # Generate a Tailwind-styled HTML summary with dark mode support via external template
        try:
                generate_lead_html(row=row, company_name=company_name, target_dir=target_dir, mapping=placeholders, overwrite=overwrite)
        except Exception as e:
                logging.error(f"Failed generating HTML summary for {company_name}: {e}")

        return output_doc


def guess_company_column(df: pd.DataFrame) -> str:
        candidates = ["Company", "Firma", "Unternehmen", "Name", "Company Name"]
        lower_map = {c.lower(): c for c in df.columns}
        for cand in candidates:
                if cand.lower() in lower_map:
                        return lower_map[cand.lower()]
                        
        # Fallback: first column
        return df.columns[0]


# ---------------- Main Pipeline ---------------- #

def main():
        parser = argparse.ArgumentParser(description="Generate website offer sheets for leads.")
        parser.add_argument("-e", "--excel", default=DEFAULT_EXCEL, help="Path to Excel file.")
        parser.add_argument("-t", "--template", default=DEFAULT_TEMPLATE, help="Path to DOCX template.")
        parser.add_argument("-o", "--output", default=DEFAULT_OUTPUT_DIR, help="Output root directory.")
        parser.add_argument("--overwrite", action="store_true", help="Overwrite existing generated offers.")
        parser.add_argument("-v", "--verbose", action="store_true", help="Verbose logging.")
        args = parser.parse_args()

        setup_logging(args.verbose)

        # Load environment variables from a .env file if present
        try:
                load_dotenv()
                logging.debug("Loaded environment from .env (if present)")
        except Exception as e:
                logging.debug(f"Failed loading .env: {e}")

        excel_path = Path(args.excel)
        template_path = Path(args.template)
        output_root = Path(args.output)

        if not excel_path.is_file():
                logging.error(f"Excel file not found: {excel_path}")
                sys.exit(1)
        if not template_path.is_file():
                logging.error(f"Template DOCX not found: {template_path}")
                sys.exit(1)

        df = load_excel(excel_path)
        if df.empty:
                logging.info("No rows found in Excel file; nothing to do.")
                return

        # Determine columns
        try:
                website_col = find_website_column(df)
        except Exception as e:
                logging.error(f"Could not determine website column: {e}")
                sys.exit(1)

        company_col = guess_company_column(df)

        filtered = filter_rows(df, website_col)
        if filtered.empty:
                logging.info("No leads matched the filter (empty website or facebook). Nothing to generate.")
                return

        ensure_dir(output_root)

        total = len(filtered)
        logging.info(f"Generating offers for {total} leads...")

        for idx, row in filtered.iterrows():
                try:
                        generate_offer(row=row, template_path=template_path, company_col=company_col, output_root=output_root, overwrite=args.overwrite)
                except Exception as e:
                        logging.error(f"Failed processing row {idx}: {e}")


if __name__ == "__main__":
        main()